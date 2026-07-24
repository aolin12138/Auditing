"""Generate mid-year report per UoA P4P template guidelines + assignment spec."""
from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup: A4, margins per template ──
for section in doc.sections:
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Mm(25); section.bottom_margin = Mm(25)
    section.left_margin = Mm(20); section.right_margin = Mm(20)
    # Page numbers in footer, centered (simple approach)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

# ── Styles ──
# Normal: Times Roman 11pt, fully justified
style = doc.styles['Normal']
font = style.font; font.name = 'Times New Roman'; font.size = Pt(11)
pf = style.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(6); pf.space_before = Pt(0)
pf.line_spacing = 1.15

# Heading 1: 12pt bold, ALL CAPS, left aligned
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'; h1.font.size = Pt(12); h1.font.bold = True
h1.font.all_caps = True; h1.font.color.rgb = RGBColor(0,0,0)
h1pf = h1.paragraph_format; h1pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1pf.space_before = Pt(12); h1pf.space_after = Pt(6)

# Heading 2: 11pt bold, Sentence case, left aligned
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'; h2.font.size = Pt(11); h2.font.bold = True
h2.font.color.rgb = RGBColor(0,0,0)
h2pf = h2.paragraph_format; h2pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2pf.space_before = Pt(8); h2pf.space_after = Pt(4)

# Heading 3: 11pt italic, not bold, Sentence case
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'; h3.font.size = Pt(11); h3.font.bold = False
h3.font.italic = True; h3.font.color.rgb = RGBColor(0,0,0)
h3pf = h3.paragraph_format; h3pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3pf.space_before = Pt(6); h3pf.space_after = Pt(3)

FIG_DIR = 'figures/report'

def fig(filename, caption, w=Inches(4.8)):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f'[Figure: {filename}]')
        return
    # Centered figure
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=w)
    # Caption: centered, 10pt, below figure
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption); r.font.size = Pt(10); r.font.italic = True
    r.font.name = 'Times New Roman'

def add_para(text):
    """Add a justified paragraph."""
    p = doc.add_paragraph(text)
    return p


# ═══════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'When a training dataset contains structural defects, such as under-sampled '
    'subpopulations or systematic gaps in feature coverage, the resulting machine '
    'learning model may perform deceptively well on held-out test data while encoding '
    'hidden biases. This report investigates whether the spatial geometry of adversarial '
    'examples can reveal such defects in a black-box setting, without access to the '
    'training data or the data-generating process. Prior work measured the density of '
    'adversarial points generated from a decision tree under a coverage-gap defect '
    'and found decreasing density with increasing bias. We extend this approach in '
    'three directions: cross-model validation using an RBF-kernel SVM alongside the '
    'decision tree, a decision-based black-box attack (HopSkipJump) that runs identically '
    'on both models, and testing of a second defect type (label noise) to establish '
    'whether the geometric signal is specific to structured data defects or generalises '
    'to any form of data corruption.'
)


# ═══════════════════════════════════════
# 2. EXPERIMENTAL SETUP
# ═══════════════════════════════════════
doc.add_heading('2. Experimental setup', level=1)

doc.add_heading('2.1 Pipeline', level=2)
doc.add_paragraph(
    'All experiments use the iris dataset (150 samples, 4 continuous features, 3 balanced '
    'classes) and follow a five-stage pipeline. First, a defect is injected into the '
    'dataset. Second, 5-fold stratified cross-validation is performed; the model is trained '
    'on each training fold. Third, an adversarial attack generates perturbations on the '
    'correctly-classified test-fold points; only successful perturbations (those that change '
    'the predicted label) are retained. Fourth, OPTICS clustering groups the adversarial '
    'points. Fifth, per cluster, three quantities are computed: the number of points, the '
    'mean pairwise Euclidean distance (geometric spread), and the cluster density defined '
    'as n_pairs divided by the sum of pairwise distances. This density formulation is '
    'approximately proportional to 1 / mean_pairwise_distance — the inverse of geometric '
    'spread — and is independent of cluster size. Metrics are averaged across clusters, '
    'then across the five folds.'
)

doc.add_heading('2.2 Defects', level=2)
doc.add_paragraph(
    'Two defect types are tested. Coverage gap: for a target class, points are sorted '
    'along one feature, and the bottom fraction (bias = 0.1, 0.3, 0.5, 0.7, 0.9) is '
    'deleted from the dataset, simulating a systematically under-sampled subpopulation. '
    'Label noise: a fraction (noise = 0.1, 0.2, 0.3, 0.4, 0.5) of training labels are '
    'randomly flipped to a different class; test labels remain clean, simulating '
    'annotation errors.'
)

doc.add_heading('2.3 Models and attacks', level=2)
doc.add_paragraph(
    'Two classifiers are used. DecisionTreeClassifier (max_depth = 3) provides continuity '
    'with prior work. SVC with RBF kernel (probability = True) handles multi-class '
    'classification via one-vs-one (three pairwise SVMs for iris), representing a '
    'fundamentally different hypothesis class with a smooth rather than axis-aligned '
    'decision boundary. Two adversarial attacks are used. DecisionTreeAttack is a '
    'white-box attack that walks the tree structure to find the minimal perturbation '
    'flipping a leaf assignment; it is fast and exact but tree-specific, used as a '
    'baseline. HopSkipJump (HSJ) is a decision-based black-box attack requiring only '
    'predicted class labels: it starts from a correctly classified point, jumps to a '
    'point of another class, performs binary search to locate the decision boundary, '
    'and iteratively refines using estimated gradients. Parameters: L2 norm, max_iter = 10, '
    'max_eval = 200, init_eval = 50, untargeted. Because HSJ needs no model internals, '
    'the identical configuration runs on both tree and SVM, enabling cross-model comparison.'
)

doc.add_heading('2.4 OPTICS clustering', level=2)
doc.add_paragraph(
    'OPTICS (Ordering Points To Identify the Clustering Structure) is used rather than '
    'DBSCAN. DBSCAN requires a global density threshold (epsilon) — all clusters must '
    'share the same density, so it will either merge clusters of different densities '
    'or discard sparse clusters as noise. Since these experiments measure changes in '
    'adversarial-point density and spread, the clustering method must be density-adaptive. '
    'OPTICS constructs a reachability plot and extracts clusters at varying density levels '
    'via the xi parameter without pre-filtering by a global threshold. Parameters: '
    'min_samples = 3, xi = 0.05, min_cluster_size = 3.'
)

doc.add_heading('2.5 Multi-dataset extension', level=2)
doc.add_paragraph(
    'In addition to iris, label noise was tested on wine (well-separated continuous '
    'features) and Car Evaluation (categorical features), and on three synthetic 3D '
    'datasets designed to isolate the effect of class separability: well-separated '
    '(class_sep = 1.5), overlapping (class_sep = 0.6, two subclusters per class), and '
    'categorical (features discretised into 4-5 bins). These used a DecisionTreeClassifier '
    'with DecisionTreeAttack.'
)

doc.add_heading('2.6 Experiment grids', level=2)
table = doc.add_table(rows=4, cols=5, style='Table Grid')
for i, h in enumerate(['Grid', 'Models', 'Attack', 'Defects', 'Runs']):
    cell = table.rows[0].cells[i]; cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True; cell.paragraphs[0].runs[0].font.size = Pt(10)
for r, d in enumerate([
    ['1', 'DecisionTree', 'DecisionTreeAttack', 'Coverage gap + Label noise', '360'],
    ['2', 'DecisionTree + SVC', 'HopSkipJump', 'Coverage gap', '360'],
    ['3', 'DecisionTree + SVC', 'HopSkipJump', 'Label noise', '360'],
], 1):
    for c, v in enumerate(d):
        cell = table.rows[r].cells[c]; cell.text = v
        cell.paragraphs[0].runs[0].font.size = Pt(10)
# Table caption above (per template)
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(4)
r = cap.add_run('Table 1. Experiment grids. Each row runs 5-fold CV on iris, 3 seeds, all class\u00d7feature combinations.')
r.font.size = Pt(10); r.font.italic = True; r.font.name = 'Times New Roman'


# ═══════════════════════════════════════
# 3. RESULTS
# ═══════════════════════════════════════
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Label noise', level=2)
doc.add_paragraph(
    'Label noise was tested on three real datasets using DecisionTree + DecisionTreeAttack '
    'and on both models (tree and SVM) using the HopSkipJump black-box attack.'
)
doc.add_paragraph(
    'On iris, adversarial-point spread (mean pairwise distance) increased steadily as label '
    'noise rose from 0.1 to 0.5 for all three model-attack combinations, with large effect '
    'sizes (Cohen\'s d = +2.0 for Tree + DecisionTreeAttack, +1.6 for Tree + HopSkipJump, '
    'and +1.1 for SVM + HopSkipJump; Figure 1). The direction of this effect, however, '
    'depends on class separability rather than on defect severity alone. Repeating the '
    'measurement on wine (well-separated continuous), Car Evaluation (categorical, highly '
    'overlapping), and three synthetic 3D datasets varying only in class separability showed '
    'that well-separated classes produce increasing spread while overlapping classes produce '
    'little change or the opposite trend. The mechanism: when classes are well-separated, '
    'label noise drives the tree to add leaves that memorise corrupted labels, fragmenting '
    'the decision boundary and scattering adversarial points; when classes already overlap, '
    'the tree is near its representational ceiling, so additional noise creates few new '
    'boundary regions and adversarial points remain clustered in the same confusion zones.'
)

fig('p2_label_noise.png',
    'Figure 1. Label noise across the full 0.1\u20130.9 range for all three model-attack '
    'combinations. Top: mean pairwise distance with 95% CI error bars. Bottom: test accuracy; '
    'the dotted line marks the 1/3 chance level for the three-class problem. Spread rises with '
    'noise but test accuracy falls in lockstep. Beyond noise 0.5 (shaded region) the model '
    'approaches a random classifier and the metric destabilises: the variance grows several-'
    'fold and many runs fail to yield clusterable adversarial sets.')

doc.add_paragraph(
    'Two observations make label noise a negative result. First, the geometric signal is '
    'confounded with accuracy: across the stable 0.1\u20130.5 range, test accuracy declines '
    'in step with the rising spread (tree 0.93 to 0.64; SVM 0.96 to 0.78). Because the defect '
    'is already plainly visible in test accuracy, the geometry adds no diagnostic information. '
    'Second, beyond noise 0.5 the metric becomes unusable. As the fraction of flipped labels '
    'approaches and exceeds one half, the trained model tends toward a random classifier '
    '(accuracy falls below the 1/3 chance level by noise 0.8), the decision surface is '
    'saturated with spurious boundaries, and both the set of still-correctly-classified test '
    'points and the landing sites of the stochastic HopSkipJump attack become dominated by '
    'randomness. The measured spread variance grows by roughly an order of magnitude and the '
    'number of valid runs collapses, so no stable trend exists. Label noise therefore '
    'provides no independent geometric diagnostic in either regime.'
)

doc.add_heading('3.2 Coverage gap', level=2)
doc.add_paragraph(
    'Coverage gap produced a different result. Figure 2 shows the geometric spread '
    '(mean pairwise distance) of adversarial points plotted against bias level across '
    'three model-attack combinations.'
)

fig('p1_coverage_gap.png',
    'Figure 2. Coverage gap: geometric spread increases with bias. Top panel shows mean '
    'pairwise distance with 95% CI error bars. Bottom panel shows test accuracy, which '
    'remains flat at approximately 0.96 across all bias levels.')

doc.add_paragraph(
    'Geometric spread increased monotonically with bias on all three combinations: '
    'Tree + DecisionTreeAttack (Cohen\'s d = +2.06), SVM + HopSkipJump (d = +0.75), '
    'and Tree + HopSkipJump (d = +0.38). The Tree + HSJ combination produced the '
    'weakest signal because HopSkipJump estimates the direction to the boundary by '
    'sampling perturbations and observing which side of the boundary they fall on. A '
    'decision tree\'s boundary is a set of flat, axis-aligned facets; while the normal '
    'direction is well-defined on each facet, most perturbations remain inside the same '
    'leaf region without crossing the boundary, so the estimated direction is noisy and '
    'the attack frequently fails to converge (many such runs had to be terminated). '
    'Critically, test accuracy remained flat across all '
    'bias levels at approximately 0.96: the model classified test points correctly '
    'regardless of the coverage gap, yet the adversarial geometry exposed the defect.'
)

doc.add_heading('3.3 Verification of the signal', level=2)
doc.add_paragraph(
    'A potential confound is that the coverage gap is injected before the cross-validation '
    'split, so the test-set composition changes with bias. An increase in geometric spread '
    'could simply reflect the changing test data rather than a genuine attack effect. '
    'To test this, two quantities were measured (Figure 3). First, the compression ratio: '
    'adversarial-point spread divided by the spread of the original correctly-classified '
    'test points from which they were generated. A ratio approaching 1.0 indicates the '
    'adversarial cloud merely mirrors the original test-point geometry. Second, the mean '
    'L2 perturbation magnitude: how far adversarial points moved from their starting '
    'positions.'
)

fig('p4_discriminant.png',
    'Figure 3. Left: compression ratio (adv_spread / orig_spread). Right: perturbation '
    'magnitude (L2). Red = label noise (Tree + DTA); green = coverage gap (SVM + HSJ). '
    'Label noise converges to 1.0 and perturbation decreases; coverage gap stays at ~0.70 '
    'and perturbation increases.')

doc.add_paragraph(
    'Under label noise, the compression ratio rose from 0.77 at noise level 0 to 0.98 '
    'at noise level 0.5: the adversarial cloud converged to the original test-point '
    'geometry, and the perturbation magnitude shrank from 0.81 to 0.64. The geometric '
    'signal under label noise was measuring the fixed test-data distribution, not the '
    'attack. Under coverage gap, the compression ratio remained at approximately 0.70 '
    'across all bias levels: the adversarial cloud was meaningfully compressed — about '
    '30% tighter than the original test points. The perturbation magnitude increased '
    'from 1.38 to 1.70, meaning adversarial points had to travel farther to cross the '
    'boundary as the coverage gap widened. The coverage-gap geometric signal is a genuine '
    'consequence of the attack, not an artifact of the test-set distribution.'
)


# ═══════════════════════════════════════
# 4. DISCUSSION
# ═══════════════════════════════════════
doc.add_heading('4. Discussion', level=1)

doc.add_paragraph(
    'The central finding is that coverage-gap bias in a training dataset is detectable '
    'from adversarial-example geometry in a setting where test accuracy reveals nothing. '
    'The geometric spread of adversarial points increased with bias on both a decision '
    'tree (d = +2.06 with DTA, +0.38 with HSJ) and an RBF-kernel SVM (d = +0.75 with '
    'HSJ), while test accuracy remained unchanged. The compression-ratio analysis '
    'confirms the signal is not an artifact of the test-set distribution changing with '
    'bias: the adversarial cloud is genuinely compressed to approximately 70% of the '
    'original test-point spread, and the perturbation magnitude grows with the gap.'
)

doc.add_paragraph(
    'The likely mechanism is that when a contiguous region of a class is deleted from '
    'training, the model places its decision boundary in an extrapolated, data-free zone. '
    'Adversarial perturbations that successfully cross this boundary scatter rather than '
    'concentrating at natural weak spots, since there is no training-data density in the '
    'gap region to constrain them. The surviving training data is still sufficient for '
    'accurate classification (test accuracy remains high), but the boundary placement '
    'is geometrically exposed by the attack.'
)

doc.add_paragraph(
    'Label noise does not provide a comparable diagnostic. First, when noise is severe '
    'enough to shift the adversarial geometry, test accuracy has already collapsed: the '
    'defect is visible without adversarial analysis. Second, the direction of the '
    'geometric change depends on class separability: well-separated data shows increasing '
    'spread (more boundaries cause points to scatter), while overlapping or categorical '
    'data shows little change or the opposite (few new boundary regions are created). '
    'Beyond a noise fraction of 0.5 the model approaches a random classifier and the metric '
    'destabilises entirely (variance grows by roughly an order of magnitude). This data-'
    'dependence makes label-noise geometry unreliable as a general diagnostic. Third, '
    'the compression-ratio analysis shows the label-noise signal is not measuring attack '
    'behaviour but rather the fixed original test-point distribution.'
)

doc.add_paragraph(
    'Several limitations bound these findings. The experiments use iris, a low-dimensional '
    'dataset (4 features, 150 samples). OPTICS clustering is known to degrade in higher '
    'dimensions as reachability distances concentrate; whether the geometric signal '
    'survives in more realistic feature spaces requires validation. The Tree + HSJ '
    'signal was weak (d = +0.38) because the tree\'s flat, axis-aligned decision '
    'boundary yields a noisy sampling-based direction estimate, causing HSJ to '
    'frequently fail to converge. The coverage gap was injected before the CV split, '
    'so test-set composition varies with bias: although the compression-ratio analysis '
    'partially addresses this, a controlled experiment with bias injected into training '
    'folds only would be more rigorous. The multi-dataset label-noise comparison used an '
    'earlier formulation of the density metric and did not record geometric spread '
    'separately. Future work should validate the coverage-gap signal in higher-dimensional '
    'settings, extend to additional defect types (feature noise, class imbalance, '
    'outliers), and test with neural-network classifiers using gradient-based attacks.'
)

doc.add_paragraph(
    'These findings suggest a discrimination capability: adversarial geometry can '
    'distinguish between structured data defects that warp the training distribution '
    'without corrupting the model\'s accuracy (coverage gaps) and defects that corrupt '
    'the decision boundary directly (label noise). If this holds in higher dimensions '
    'and across more model families, adversarial geometry could serve as a component '
    'of data-quality auditing pipelines, flagging datasets that appear healthy to '
    'accuracy-based evaluation but carry hidden structural biases.'
)


# ═══════════════════════════════════════
# 5. CONCLUSION
# ═══════════════════════════════════════
doc.add_heading('5. Conclusion', level=1)
doc.add_paragraph(
    'Adversarial-example geometry can detect coverage-gap bias in training data that '
    'test accuracy alone cannot see. The geometric spread of adversarial points increases '
    'with the severity of the coverage gap across two model families without any '
    'accompanying decline in test accuracy. Label noise produces no independent geometric '
    'signal: its effects are either visible through accuracy decline or depend on the '
    'separability of the original data distribution. The approach shows promise as a '
    'black-box diagnostic for structured data defects, pending validation in higher '
    'dimensions and across additional defect types.'
)


# ── Save ──
out = 'midyear_report.docx'
doc.save(out)
print(f'Saved {out}')
